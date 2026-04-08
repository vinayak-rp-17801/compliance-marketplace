import os
import re
import zipfile
import tempfile
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches

# Input files
DOCS_ZIP = "processed-documents.zip"
SCREENSHOT_ZIPS = [
    "Screenshot marketplace 1.zip",
    "Screenshot market place 2.zip",
    "Screenshot marketplace 3.zip",
    "Screenshot marketplace 4.zip",
]
OUTPUT_ZIP = "output-documents.zip"

# Screenshots to add in order
SCREENSHOT_NAMES = [
    "01_overview.png",
    "02_graph_overview.png",
    "03_report_detail.png",
    "04_manage_compliance.png",
]

# Screenshot dimensions at 96 DPI
SCREENSHOT_WIDTH_INCHES = 7.29   # ~700px wide, fits within typical page margins
SCREENSHOT_HEIGHT_INCHES = 4.10  # proportional to 1366x768


def extract_compliance_name(doc_filename):
    """Extract compliance name from a filename like '001_201 CMR 17.00.docx'."""
    name = re.sub(r"^\d+_", "", doc_filename)
    name = name.replace(".docx", "")
    return name


def build_screenshot_map(screenshot_zips):
    """
    Build a mapping of compliance_name -> {screenshot_name: zip_internal_path}.
    Also returns a dict of zip_name -> ZipFile for later reading.
    """
    screenshot_map = {}  # compliance_name -> {screenshot_name: (zip_name, internal_path)}

    for zip_name in screenshot_zips:
        if not os.path.exists(zip_name):
            print(f"  WARNING: {zip_name} not found, skipping.")
            continue
        with zipfile.ZipFile(zip_name, "r") as z:
            for item in z.namelist():
                # Skip Mac metadata files and directories
                if "__MACOSX" in item or item.split("/")[-1].startswith("._"):
                    continue
                parts = item.strip("/").split("/")
                # Expected structure: <main_folder>/<compliance_name>/<screenshot.png>
                if len(parts) != 3:
                    continue
                compliance_name = parts[1]
                screenshot_file = parts[2]
                if screenshot_file in SCREENSHOT_NAMES:
                    if compliance_name not in screenshot_map:
                        screenshot_map[compliance_name] = {}
                    screenshot_map[compliance_name][screenshot_file] = (zip_name, item)

    return screenshot_map


METADATA_LINES = [
    "Valid Submission : Yes",
    "Supported Editions : Free, Basic, Standard, Professional, MSSP",
    "Help Document : https://www.manageengine.com/log-management/help/integrations/compliance-extensions.html",
    "Help Video : Not provided",
    "Tags :  Log360 Cloud, Compliance, Privacy monitoring, Data privacy, Regulatory compliance",
    "Supported DCs : US, IN, JP, CA, AU, UK and EU",
    "Privacy Policy:",
    "Terms of Service:",
    "Release Notes:",
]


def add_screenshots_to_doc(doc_path, screenshots, output_path, compliance_name=""):
    """
    Load a DOCX document, append 4 screenshots in order, add metadata content,
    and save to output_path.
    screenshots: list of local file paths in the correct order.
    compliance_name: name extracted from the document filename, used in the last metadata line.
    """
    doc = Document(doc_path)

    for screenshot_path in screenshots:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(
            screenshot_path,
            width=Inches(SCREENSHOT_WIDTH_INCHES),
        )
        paragraph.alignment = 1  # Center

    # Add metadata content after screenshots
    for line in METADATA_LINES:
        doc.add_paragraph(line)

    # Last line with dynamic compliance name
    predefined_line = f"Predefined reports for the {compliance_name}."
    doc.add_paragraph(predefined_line)

    doc.save(output_path)


def main():
    print("=" * 70)
    print("STEP 1: Building screenshot map from screenshot zip files...")
    print("=" * 70)

    screenshot_map = build_screenshot_map(SCREENSHOT_ZIPS)
    print(f"Found screenshots for {len(screenshot_map)} compliance folders.\n")

    print("=" * 70)
    print("STEP 2: Processing compliance documents from processed-documents.zip...")
    print("=" * 70 + "\n")

    with tempfile.TemporaryDirectory() as tmp_dir:
        docs_extract = os.path.join(tmp_dir, "docs")
        screenshots_extract = os.path.join(tmp_dir, "screenshots")
        output_docs_dir = os.path.join(tmp_dir, "output_docs")

        os.makedirs(docs_extract, exist_ok=True)
        os.makedirs(screenshots_extract, exist_ok=True)
        os.makedirs(output_docs_dir, exist_ok=True)

        # Extract screenshots from all zip files
        for zip_name in SCREENSHOT_ZIPS:
            if not os.path.exists(zip_name):
                continue
            with zipfile.ZipFile(zip_name, "r") as z:
                for item in z.namelist():
                    if "__MACOSX" in item or item.split("/")[-1].startswith("._"):
                        continue
                    parts = item.strip("/").split("/")
                    if len(parts) == 3 and parts[2] in SCREENSHOT_NAMES:
                        target_path = os.path.join(
                            screenshots_extract, parts[1], parts[2]
                        )
                        os.makedirs(os.path.dirname(target_path), exist_ok=True)
                        with z.open(item) as src, open(target_path, "wb") as dst:
                            shutil.copyfileobj(src, dst)

        # Extract compliance documents
        with zipfile.ZipFile(DOCS_ZIP, "r") as z:
            for item in z.namelist():
                if item.endswith(".docx") and not item.split("/")[-1].startswith("._"):
                    target_path = os.path.join(docs_extract, item.split("/")[-1])
                    with z.open(item) as src, open(target_path, "wb") as dst:
                        shutil.copyfileobj(src, dst)

        # Process each document
        doc_files = sorted(
            [f for f in os.listdir(docs_extract) if f.endswith(".docx")]
        )
        total = len(doc_files)
        successful = 0
        failed = 0
        no_screenshots = 0

        for idx, doc_file in enumerate(doc_files, 1):
            compliance_name = extract_compliance_name(doc_file)
            doc_path = os.path.join(docs_extract, doc_file)
            output_path = os.path.join(output_docs_dir, doc_file)

            print(f"[{idx:3d}/{total}] {doc_file}")
            print(f"         Compliance: {compliance_name}")

            if compliance_name not in screenshot_map:
                print(f"         WARNING: No screenshot folder found for '{compliance_name}', skipping screenshots.\n")
                # Copy document as-is
                shutil.copy2(doc_path, output_path)
                no_screenshots += 1
                continue

            screenshot_paths = []

            for screenshot_name in SCREENSHOT_NAMES:
                sc_path = os.path.join(screenshots_extract, compliance_name, screenshot_name)
                if os.path.exists(sc_path):
                    screenshot_paths.append(sc_path)
                    print(f"         ✓ {screenshot_name}")
                else:
                    print(f"         ✗ {screenshot_name} not found")

            try:
                add_screenshots_to_doc(doc_path, screenshot_paths, output_path, compliance_name)
                print(f"         ✓ Document saved\n")
                successful += 1
            except Exception as e:
                print(f"         ✗ Error saving document: {e}\n")
                shutil.copy2(doc_path, output_path)
                failed += 1

        # Package output documents into a zip file
        print("=" * 70)
        print("STEP 3: Creating output zip file...")
        print("=" * 70)

        output_files = sorted(os.listdir(output_docs_dir))
        with zipfile.ZipFile(OUTPUT_ZIP, "w", zipfile.ZIP_DEFLATED) as zout:
            for doc_file in output_files:
                file_path = os.path.join(output_docs_dir, doc_file)
                zout.write(file_path, arcname=doc_file)

        print(f"Output written to: {OUTPUT_ZIP}")
        print(f"Total documents in output: {len(output_files)}\n")

    # Summary
    print("=" * 70)
    print("PROCESSING COMPLETE")
    print("=" * 70)
    print(f"Total documents:              {total}")
    print(f"Successfully added screenshots: {successful}")
    print(f"No screenshots found:           {no_screenshots}")
    print(f"Errors:                         {failed}")
    print(f"Output zip:                     {OUTPUT_ZIP}")
    print("=" * 70)


if __name__ == "__main__":
    main()
