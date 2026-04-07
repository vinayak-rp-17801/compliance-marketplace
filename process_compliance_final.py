import os
import zipfile
from pathlib import Path
from docx import Document
from docx.shared import Inches
import re

def sanitize_folder_name(doc_name):
    """Convert '001_201 CMR 17.00.docx' to '201_CMR_17_00'"""
    # Remove leading number and underscore (001_, 002_, etc.)
    name = re.sub(r'^\d+_', '', doc_name)
    # Remove .docx extension
    name = name.replace('.docx', '')
    # Replace spaces with underscores
    name = name.replace(' ', '_')
    # Replace dots with underscores
    name = name.replace('.', '_')
    return name

def extract_zips(docs_zip, logos_zip, extract_dir):
    docs_extract = os.path.join(extract_dir, 'docs')
    logos_extract = os.path.join(extract_dir, 'logos')
    
    os.makedirs(docs_extract, exist_ok=True)
    os.makedirs(logos_extract, exist_ok=True)
    
    print("Extracting compliance_docs.zip...")
    with zipfile.ZipFile(docs_zip, 'r') as zip_ref:
        zip_ref.extractall(docs_extract)
    
    print("Extracting compliance logo marketplace.zip...")
    with zipfile.ZipFile(logos_zip, 'r') as zip_ref:
        zip_ref.extractall(logos_extract)
    
    return docs_extract, logos_extract

def find_images(logos_extract, folder_name):
    """Find logo and thumbnail images for a compliance"""
    # Try looking in the subdirectory first (compliance logo marketplace/)
    folder_path = os.path.join(logos_extract, 'compliance logo marketplace', folder_name)
    
    if not os.path.exists(folder_path):
        # Fallback to direct path (in case structure is different)
        folder_path = os.path.join(logos_extract, folder_name)
    
    if not os.path.exists(folder_path):
        return None, None
    
    logo_path = None
    thumbnail_path = None
    
    try:
        for file in os.listdir(folder_path):
            if '180x180' in file or '180_180' in file:
                logo_path = os.path.join(folder_path, file)
            elif '740x340' in file or '740_340' in file:
                thumbnail_path = os.path.join(folder_path, file)
    except:
        pass
    
    return logo_path, thumbnail_path
def add_images_to_document(doc_path, logo_path, thumbnail_path, output_path):
    """Add logo and thumbnail images to a DOCX document"""
    try:
        doc = Document(doc_path)
        
        # Add logo (180x180) - 1.875 inches at 96 DPI
        if logo_path and os.path.exists(logo_path):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.875), height=Inches(1.875))
            paragraph.alignment = 1  # Center
        
        # Add thumbnail (740x340) - 7.708 x 3.542 inches at 96 DPI
        if thumbnail_path and os.path.exists(thumbnail_path):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(thumbnail_path, width=Inches(7.708), height=Inches(3.542))
            paragraph.alignment = 1  # Center
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"  Error: {str(e)}")
        return False

def main():
    docs_zip = 'compliance_docs.zip'
    logos_zip = 'compliance logo marketplace.zip'
    extract_dir = 'extracted_files'
    output_dir = 'output_docs'
    
    os.makedirs(output_dir, exist_ok=True)
    
    print("=" * 70)
    print("STEP 1: Extracting zip files...")
    print("=" * 70)
    docs_extract, logos_extract = extract_zips(docs_zip, logos_zip, extract_dir)
    
    print("\n" + "=" * 70)
    print("STEP 2: Processing compliance documents...")
    print("=" * 70 + "\n")
    
    # Get all DOCX files and filter out Mac system files
    all_files = list(Path(docs_extract).rglob("*.docx"))
    docx_files = sorted([f for f in all_files if not f.name.startswith('._')])
    
    total_files = len(docx_files)
    successful = 0
    failed = 0
    
    print(f"Found {total_files} DOCX files\n")
    
    for idx, doc_path in enumerate(docx_files, 1):
        doc_name = doc_path.name
        folder_name = sanitize_folder_name(doc_name)
        
        print(f"[{idx}/{total_files}] {doc_name}")
        print(f"         Looking for folder: {folder_name}")
        
        # Find images
        logo_path, thumbnail_path = find_images(logos_extract, folder_name)
        
        if logo_path:
            print(f"         ✓ Found logo")
        else:
            print(f"         ✗ Logo not found")
        
        if thumbnail_path:
            print(f"         ✓ Found thumbnail")
        else:
            print(f"         ✗ Thumbnail not found")
        
        # Create output document
        output_path = os.path.join(output_dir, doc_name)
        if add_images_to_document(str(doc_path), logo_path, thumbnail_path, output_path):
            print(f"         ✓ Document created\n")
            successful += 1
        else:
            print(f"         ✗ Failed to create document\n")
            failed += 1
    
    # Summary
    print("=" * 70)
    print("PROCESSING COMPLETE")
    print("=" * 70)
    print(f"Total documents: {total_files}")
    print(f"Successfully processed: {successful}")
    print(f"Failed: {failed}")
    print(f"Output directory: {output_dir}")
    print("=" * 70)

if __name__ == "__main__":
    main()
