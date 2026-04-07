import os
import zipfile
from docx import Document
from PIL import Image

def extract_zip_files(zip_folder):
    for item in os.listdir(zip_folder):
        if item.endswith('.zip'):
            with zipfile.ZipFile(os.path.join(zip_folder, item), 'r') as zip_ref:
                zip_ref.extractall(zip_folder)

def process_documents(doc_folder, logo_folder, thumb_folder):
    output_folder = 'output_docs'
    os.makedirs(output_folder, exist_ok=True)

    for doc_file in os.listdir(doc_folder):
        if doc_file.endswith('.docx'):
            doc = Document(os.path.join(doc_folder, doc_file))
            # Assume logo and thumbnail are in matching folders
            logo_path = os.path.join(logo_folder, os.path.splitext(doc_file)[0] + '.png')
            thumb_path = os.path.join(thumb_folder, os.path.splitext(doc_file)[0] + '.jpg')

            # Add logo
            if os.path.exists(logo_path):
                logo = Image.open(logo_path)
                logo = logo.resize((180, 180))
                doc.add_picture(logo_path, width=docx.shared.Inches(1))

            # Add thumbnail
            if os.path.exists(thumb_path):
                thumb = Image.open(thumb_path)
                thumb = thumb.resize((740, 340))
                doc.add_picture(thumb_path, width=docx.shared.Inches(1))

            output_file = os.path.join(output_folder, doc_file)
            doc.save(output_file)

if __name__ == '__main__':
    extract_zip_files('your_zip_folder')
    process_documents('your_doc_folder', 'your_logo_folder', 'your_thumb_folder')
