import zipfile
import os
from PIL import Image
from docx import Document
from docx.shared import Inches

# Define file paths
compliance_docs_zip = 'compliance_docs.zip'
compliance_logos_zip = 'compliance_logo_marketplace.zip'
output_folder = 'output_docs'

# Create output folder
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Step 1: Extract zip files
def extract_zip(zip_path, extract_to):
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        zip_ref.extractall(extract_to)

extract_zip(compliance_docs_zip, output_folder)
extract_zip(compliance_logos_zip, output_folder)

# Step 2: Match documents with logos
def match_documents_with_logos(doc_folder, logo_folder):
    # This function should implement the matching logic. Placeholder for actual logic.
    return []

matched_files = match_documents_with_logos(output_folder, output_folder)

# Step 3: Create DOCX files with images embedded
def create_docx(file_name, logo_path, thumbnail_path):
    doc = Document()
    
    # Add logo
    logo_img = Image.open(logo_path)
    logo_img = logo_img.resize((180, 180))
    logo_path_resized = os.path.join(output_folder, 'resized_' + os.path.basename(logo_path))
    logo_img.save(logo_path_resized)
    doc.add_picture(logo_path_resized, width=Inches(2))

    # Add thumbnail
    thumbnail_img = Image.open(thumbnail_path)
    thumbnail_img = thumbnail_img.resize((740, 340))
    thumbnail_path_resized = os.path.join(output_folder, 'resized_' + os.path.basename(thumbnail_path))
    thumbnail_img.save(thumbnail_path_resized)
    doc.add_picture(thumbnail_path_resized, width=Inches(4.8))

    # Save the document
    doc_path = os.path.join(output_folder, f'{file_name}.docx')
    doc.save(doc_path)

# Iterate through matched files and create DOCX files
for i, (doc_file, logo_file, thumb_file) in enumerate(matched_files):
    create_docx(f'document_{i+1}', logo_file, thumb_file)
    print(f'Created document_{i+1}.docx')

print('Process completed successfully!')